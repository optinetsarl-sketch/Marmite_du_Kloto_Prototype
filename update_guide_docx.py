import docx

path = 'DOCUMENTATION/Guide_Utilisation_Marmite_du_Kloto.docx'
doc = docx.Document(path)

# Ajouter la section des dernières fonctionnalités
p_head = doc.add_heading('Fiche de Mise à Jour — Nouveautés & Ergonomie', level=1)

doc.add_heading('1. Tri Décroissant Généralisé (Dernières Opérations / Servies en Haut)', level=2)
doc.add_paragraph(
    "Afin de faciliter la consultation rapide lors des pics d'activité :\n"
    "• Sur l'écran Ventes : les nouvelles commandes actives s'affichent immédiatement au sommet de la liste.\n"
    "• Sur l'écran Livraison : le détail des courses du jour pour chaque livreur classe les livraisons les plus récentes en 1ère position.\n"
    "• Sur l'écran Cuisine : l'historique des repas servis présente les derniers plats préparés tout en haut.\n"
    "• Sur l'écran À emporter : l'historique des ventes à emporter clôturées affiche les opérations les plus récentes en tête.\n"
    "• Sur les écrans Dépenses & Stock : le journal des dépenses et des mouvements de stock classe les entrées de la plus récente à la plus ancienne."
)

doc.add_heading('2. Émission Immédiate du Bon de Livraison (Commandes Bar / Sans Cuisine)', level=2)
doc.add_paragraph(
    "Lorsqu'une commande à livrer ne contient que des articles de bar (boissons, bières, vin...) et aucun plat de cuisine :\n"
    "• La validation génère et ouvre automatiquement le Bon de Livraison officiel (nom du client, adresse, livreur attribué, montant à encaisser).\n"
    "• Aucun bon de cuisine inutile n'est imprimé et le livreur peut partir immédiatement."
)

doc.add_heading('3. Historique Consultable par Date sur l\'Écran À Emporter', level=2)
doc.add_paragraph(
    "L'écran À emporter (/emporter) intègre désormais un sélecteur de date interactif (avec calendrier et navigation jour par jour ‹ et ›) :\n"
    "• Permet de consulter l'historique complet des commandes à emporter clôturées pour n'importe quelle date passée.\n"
    "• Propose un bouton 📄 Reçu pour chaque vente afin d'imprimer ou réimprimer le reçu officiel."
)

doc.add_heading('4. Gestion et Suivi des Comptes Livreurs', level=2)
doc.add_paragraph(
    "Sur l'écran Livraison (/livraison), tous les livreurs sont clairement répertoriés avec le détail de leurs courses du jour, leurs recettes encaissées à remettre en caisse et l'historique par date."
)

doc.save(path)
print("Docx mis à jour avec succès !")
