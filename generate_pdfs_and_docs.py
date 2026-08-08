import os
import shutil
import docx

# 1. Update DOCX files in FACTURE
doc_bl_path = 'FACTURE/Bordereau_Livraison_Marmite_du_Kloto.docx'
if os.path.exists(doc_bl_path):
    doc_bl = docx.Document(doc_bl_path)
    for p in doc_bl.paragraphs:
        if '30/07/2026' in p.text:
            p.text = p.text.replace('30/07/2026', '04/08/2026')
        if '30 Juillet 2026' in p.text:
            p.text = p.text.replace('30 Juillet 2026', '04 Août 2026')
        if '2026/07' in p.text:
            p.text = p.text.replace('2026/07', '2026/08')
    for table in doc_bl.tables:
        for row in table.rows:
            for cell in row.cells:
                if '30/07/2026' in cell.text:
                    cell.text = cell.text.replace('30/07/2026', '04/08/2026')
                if '30 Juillet 2026' in cell.text:
                    cell.text = cell.text.replace('30 Juillet 2026', '04 Août 2026')
                if '2026/07' in cell.text:
                    cell.text = cell.text.replace('2026/07', '2026/08')
    doc_bl.save(doc_bl_path)

doc_fac_path = 'FACTURE/Facture_Definitive_Marmite_du_Kloto_150000.docx'
if os.path.exists(doc_fac_path):
    doc_fac = docx.Document(doc_fac_path)
    for p in doc_fac.paragraphs:
        if '30/07/2026' in p.text:
            p.text = p.text.replace('30/07/2026', '04/08/2026')
        if '30 Juillet 2026' in p.text:
            p.text = p.text.replace('30 Juillet 2026', '04 Août 2026')
        if '2026/07' in p.text:
            p.text = p.text.replace('2026/07', '2026/08')
    for table in doc_fac.tables:
        for row in table.rows:
            for cell in row.cells:
                if '30/07/2026' in cell.text:
                    cell.text = cell.text.replace('30/07/2026', '04/08/2026')
                if '30 Juillet 2026' in cell.text:
                    cell.text = cell.text.replace('30 Juillet 2026', '04 Août 2026')
                if '2026/07' in cell.text:
                    cell.text = cell.text.replace('2026/07', '2026/08')
    doc_fac.save(doc_fac_path)

# 2. Copy into deployment_client/DOCUMENTATION and deployment_client/FACTURE and DOCUMENTATION
dest_doc = "deployment_client/DOCUMENTATION"
dest_fac = "deployment_client/FACTURE"

os.makedirs(dest_doc, exist_ok=True)
os.makedirs(dest_fac, exist_ok=True)

files_to_copy = [
    "FACTURE/Bordereau_Livraison_Marmite_du_Kloto.html",
    "FACTURE/Facture_Definitive_Marmite_du_Kloto_Exacte.html",
    "FACTURE/Bordereau_Livraison_Marmite_du_Kloto.docx",
    "FACTURE/Facture_Definitive_Marmite_du_Kloto_150000.docx"
]

for src in files_to_copy:
    if os.path.exists(src):
        shutil.copy(src, dest_doc)
        shutil.copy(src, dest_fac)
        shutil.copy(src, "DOCUMENTATION")
        print(f"[OK] Copié {src}")

print("Mise à jour bordereau et facture terminée !")
